"""
Convert docs/food-programme-rp-sop.md to docx.

Hand-rolled parser for the specific markdown features used in this doc:
- Headings (#, ##, ###, ####)
- Paragraphs with inline **bold**, *italic*, `code`, [text](url)
- Bullet lists (- ...)
- Checkbox lists (- [ ] ...) rendered as ☐ ...
- Tables (pipe syntax with header row + delimiter row)
- Fenced code blocks (``` ... ```)
- Blockquotes (> ...)
- Horizontal rules (---)

Run: python3 scripts/md-to-docx.py [INPUT.md] [OUTPUT.docx]
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DEFAULT_IN  = "docs/food-programme-rp-sop.md"
DEFAULT_OUT = "docs/food-programme-rp-sop.docx"

# ── helpers ────────────────────────────────────────────────────────────────

def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_hr(doc: Document) -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


INLINE_TOKEN_RE = re.compile(
    r"(\*\*[^*\n]+\*\*"     # bold
    r"|`[^`\n]+`"            # code
    r"|\*[^*\n]+\*"          # italic
    r"|\[[^\]]+\]\([^)]+\))" # link
)


def add_inline_runs(paragraph, text: str) -> None:
    """Add runs to a paragraph, interpreting **bold**, *italic*, `code`, links."""
    parts = INLINE_TOKEN_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Menlo"
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x80, 0x21, 0x21)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        elif part.startswith("[") and "](" in part and part.endswith(")"):
            text_part, url_part = part[1:-1].split("](", 1)
            r = paragraph.add_run(text_part)
            r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
            r.underline = True
        else:
            paragraph.add_run(part)


# ── table handling ─────────────────────────────────────────────────────────

def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_table_delim(line: str) -> bool:
    cells = split_table_row(line)
    if not cells:
        return False
    return all(re.fullmatch(r":?-{3,}:?", c) for c in cells)


def render_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    header, *body = rows
    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header
    for i, cell_text in enumerate(header):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "1F2937")
        p = cell.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        add_inline_runs(p, cell_text)
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)

    # body
    for r_i, row in enumerate(body):
        row = (row + [""] * len(header))[:len(header)]
        for c_i, cell_text in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            p = cell.paragraphs[0]
            add_inline_runs(p, cell_text)
            for run in p.runs:
                run.font.size = Pt(10)


# ── document build ─────────────────────────────────────────────────────────

def add_heading(doc: Document, text: str, level: int) -> None:
    sizes = {1: 22, 2: 16, 3: 13, 4: 11}
    colors = {1: 0x111827, 2: 0x111827, 3: 0x1F2937, 4: 0x374151}
    p = doc.add_paragraph()
    add_inline_runs(p, text)
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(sizes.get(level, 11))
        run.font.color.rgb = RGBColor(
            (colors[level] >> 16) & 0xFF,
            (colors[level] >> 8) & 0xFF,
            colors[level] & 0xFF,
        )
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True


def add_bullet(doc: Document, text: str, *, checkbox: bool, blockquote: bool = False) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.2 if not blockquote else 1.8)
    p.paragraph_format.space_after = Pt(2)
    if checkbox:
        r = p.add_run("☐  ")
        r.font.name = "Menlo"
    add_inline_runs(p, text)
    for run in p.runs:
        run.font.size = Pt(10.5)


def add_paragraph(doc: Document, text: str, *, blockquote: bool = False) -> None:
    p = doc.add_paragraph()
    if blockquote:
        p.paragraph_format.left_indent = Cm(0.6)
        p_pr = p._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "10")
        left.set(qn("w:color"), "9CA3AF")
        p_bdr.append(left)
        p_pr.append(p_bdr)
    add_inline_runs(p, text)
    for run in p.runs:
        run.font.size = Pt(10.5)
        if blockquote:
            run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    p.paragraph_format.space_after = Pt(6)


def add_code_block(doc: Document, lines: list[str]) -> None:
    text = "\n".join(lines)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F3F4F6")
    p_pr.append(shd)
    r = p.add_run(text)
    r.font.name = "Menlo"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x11, 0x18, 0x27)


# ── parser ─────────────────────────────────────────────────────────────────

def parse(md_text: str) -> list[tuple[str, object]]:
    """Return a list of (kind, payload) blocks."""
    lines = md_text.splitlines()
    out = []
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]

        # fenced code block
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            out.append(("code", code_lines))
            continue

        # heading
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            out.append(("heading", (len(m.group(1)), m.group(2).strip())))
            i += 1
            continue

        # horizontal rule
        if re.fullmatch(r"-{3,}|_{3,}|\*{3,}", line.strip()):
            out.append(("hr", None))
            i += 1
            continue

        # table — header line followed by delimiter line
        if line.startswith("|") and i + 1 < n and is_table_delim(lines[i + 1]):
            rows = [split_table_row(line)]
            i += 2  # skip header + delim
            while i < n and lines[i].startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            out.append(("table", rows))
            continue

        # blockquote (possibly multi-line)
        if line.startswith(">"):
            quote_lines = []
            while i < n and lines[i].startswith(">"):
                quote_lines.append(lines[i].lstrip(">").lstrip())
                i += 1
            out.append(("quote", quote_lines))
            continue

        # bullet / checkbox
        m = re.match(r"^(\s*)-\s+(.*)$", line)
        if m:
            indent_spaces = len(m.group(1))
            text = m.group(2)
            cm = re.match(r"^\[([ xX])\]\s+(.*)$", text)
            if cm:
                out.append(("checkbox", (indent_spaces, cm.group(2))))
            else:
                out.append(("bullet", (indent_spaces, text)))
            i += 1
            continue

        # blank line
        if not line.strip():
            out.append(("blank", None))
            i += 1
            continue

        # paragraph (collect until blank or block-level token)
        para_lines = [line]
        i += 1
        while i < n:
            nxt = lines[i]
            if (not nxt.strip()
                    or nxt.startswith("#")
                    or nxt.startswith("|")
                    or nxt.startswith(">")
                    or nxt.strip().startswith("```")
                    or re.match(r"^\s*-\s+", nxt)
                    or re.fullmatch(r"-{3,}|_{3,}|\*{3,}", nxt.strip())):
                break
            para_lines.append(nxt)
            i += 1
        out.append(("para", " ".join(s.strip() for s in para_lines)))

    return out


# ── main ───────────────────────────────────────────────────────────────────

def main() -> None:
    in_path = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(DEFAULT_IN)
    out_path = Path(sys.argv[2]) if len(sys.argv) > 2 else Path(DEFAULT_OUT)

    md_text = in_path.read_text(encoding="utf-8")
    blocks = parse(md_text)

    doc = Document()

    # base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # page margins
    for section in doc.sections:
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)

    for kind, payload in blocks:
        if kind == "heading":
            level, text = payload
            add_heading(doc, text, level)
        elif kind == "para":
            add_paragraph(doc, payload)
        elif kind == "bullet":
            _, text = payload
            add_bullet(doc, text, checkbox=False)
        elif kind == "checkbox":
            _, text = payload
            add_bullet(doc, text, checkbox=True)
        elif kind == "table":
            render_table(doc, payload)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        elif kind == "quote":
            for ql in payload:
                add_paragraph(doc, ql, blockquote=True)
        elif kind == "code":
            add_code_block(doc, payload)
        elif kind == "hr":
            add_hr(doc)
        elif kind == "blank":
            # already handled by spacing
            pass

    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
